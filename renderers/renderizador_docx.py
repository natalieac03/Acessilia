"""Gera o DOCX com equacoes OMML e imagens com texto alternativo.

As equacoes vao como OMML, a linguagem nativa de matematica do Word —
o arquivo abre a formula como equacao editavel, nao como imagem colada
(imagem, para leitor de tela, equivale a nada).

As imagens sao embutidas de verdade, com o texto alternativo gravado
nas propriedades da figura, que e onde o leitor de tela do Word
procura. Quando existe descricao longa, ela vai num paragrafo logo
abaixo. Imagem decorativa nao gera nada.
"""

from __future__ import annotations

from pathlib import Path
from typing import Any

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt

from pipeline.gestor_de_verbosidade import filtrar_blocos_por_perfil


def gerar_docx(document: dict[str, Any], output_path: Path, profile_name: str = "docx", filename: str = "") -> Path:
    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Arial"
    style.font.size = Pt(11)
    _definir_idioma(doc, str(document.get("language") or "pt-BR"))
    if filename:
        heading = doc.add_heading(filename, level=1)
        heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for section in document.get("sections", []):
        _render_section(doc, section, profile_name)
    doc.save(str(output_path))
    return output_path


def _render_section(doc: Document, section: dict[str, Any], profile_name: str) -> None:
    if section.get("title"):
        doc.add_heading(section["title"], level=min(section.get("level", 1), 9))
    for block in filtrar_blocos_por_perfil(section.get("blocks", []), profile_name):
        _render_block(doc, block)
    for child in section.get("children", []):
        _render_section(doc, child, profile_name)


def _render_block(doc: Document, block: dict[str, Any]) -> None:
    block_type = block.get("type")
    if block_type == "heading":
        doc.add_heading(block.get("title", block.get("text", "")), level=min(block.get("level", 1), 9))
    elif block_type == "paragraph":
        doc.add_paragraph(block.get("text", ""))
    elif block_type == "code":
        paragraph = doc.add_paragraph()
        run = paragraph.add_run(block.get("text", ""))
        run.font.name = "Courier New"
        run.font.size = Pt(10)
        paragraph.style = doc.styles["No Spacing"]
    elif block_type == "list":
        inicio = int(block.get("start") or 1)
        if block.get("ordered") and inicio != 1:
            for k, item in enumerate(block.get("items", [])):
                doc.add_paragraph(f"{inicio + k}. {item}")
            return
        style = "List Number" if block.get("ordered") else "List Bullet"
        for item in block.get("items", []):
            doc.add_paragraph(str(item), style=style)
    elif block_type == "table":
        rows = block.get("rows", [])
        if rows:
            _adicionar_caption_e_resumo(doc, block)
            table = doc.add_table(rows=len(rows), cols=max(len(row) for row in rows))
            table.style = "Table Grid"
            idiomas = block.get("cell_languages") or []
            idioma_doc = str(block.get("language") or "pt-BR")
            for i, row in enumerate(rows):
                for j, cell in enumerate(row):
                    celula = table.cell(i, j)
                    celula.text = str(cell)
                    idioma = (
                        idiomas[i][j]
                        if i < len(idiomas) and j < len(idiomas[i]) else ""
                    )
                    _marcar_idioma_da_celula(celula, idioma, idioma_doc)
            _marcar_linha_de_cabecalho(table)
    elif block_type == "image" and block.get("decorative"):
        pass
    elif block_type == "math":
        _adicionar_formula(doc, block)
    elif block_type == "image":
        _adicionar_imagem(doc, block)
    elif block_type in {"details", "note", "warning", "quote"}:
        doc.add_paragraph(block.get("text", block.get("alt_text", "")))
    else:
        doc.add_paragraph(block.get("text", ""))


_NS_MATH = "http://schemas.openxmlformats.org/officeDocument/2006/math"


def _elemento_omml(omml: str):
    if not omml or "oMath" not in omml:
        return None
    try:
        from lxml import etree
    except Exception:
        return None
    fragmento = omml.strip()
    if f'xmlns:m="{_NS_MATH}"' not in fragmento:
        fragmento = fragmento.replace(
            "<m:oMath", f'<m:oMath xmlns:m="{_NS_MATH}"', 1
        )
    try:
        return etree.fromstring(fragmento.encode("utf-8"))
    except Exception:
        return None


def _adicionar_formula(doc: Document, block: dict[str, Any]) -> None:
    from pipeline.higienizador import remover_marcadores_tecnicos

    fala = remover_marcadores_tecnicos(
        block.get("speech_pt_br") or block.get("text") or ""
    )
    elemento = _elemento_omml(block.get("omml") or "")
    if elemento is None:
        if fala:
            doc.add_paragraph(fala)
        return
    paragrafo = doc.add_paragraph()
    paragrafo._p.append(elemento)


def _marcar_linha_de_cabecalho(table) -> None:
    try:
        from docx.oxml.ns import qn

        primeira = table.rows[0]
        propriedades = primeira._tr.get_or_add_trPr()
        if propriedades.find(qn("w:tblHeader")) is None:
            propriedades.append(
                propriedades.makeelement(qn("w:tblHeader"), {})
            )
    except Exception:
        pass


def _definir_idioma(doc: Document, idioma: str) -> None:
    try:
        from docx.oxml.ns import qn

        estilo = doc.styles["Normal"]
        propriedades = estilo.element.get_or_add_rPr()
        marca = propriedades.find(qn("w:lang"))
        if marca is None:
            marca = propriedades.makeelement(qn("w:lang"), {})
            propriedades.append(marca)
        marca.set(qn("w:val"), idioma)
    except Exception:
        pass

_LARGURA_MAXIMA_POL = 6.0


def _definir_texto_alternativo(shape, alt: str, descricao: str = "") -> None:
    try:
        docPr = shape._inline.docPr
        curto = (alt or "").strip()[:255]
        docPr.set("descr", curto)
        docPr.set("title", curto)
    except Exception:
        pass


def _adicionar_imagem(doc: Document, block: dict) -> None:
    from pathlib import Path

    alt = block.get("alt_text") or block.get("text") or ""
    descricao = block.get("long_description") or ""
    caminho = block.get("asset_path") or (
        block.get("metadata", {}) or {}
    ).get("asset_path")

    inserida = False
    if caminho:
        arquivo = Path(caminho)
        if not arquivo.is_absolute():
            base = (block.get("metadata", {}) or {}).get("asset_base")
            if base:
                arquivo = Path(base) / caminho
        if arquivo.exists():
            try:
                from docx.shared import Inches

                shape = doc.add_picture(
                    str(arquivo), width=Inches(_LARGURA_MAXIMA_POL)
                )
                _definir_texto_alternativo(shape, alt, descricao)
                inserida = True
            except Exception:
                inserida = False

    if inserida and descricao:
        paragrafo = doc.add_paragraph()
        rotulo = paragrafo.add_run("Descrição detalhada: ")
        rotulo.bold = True
        paragrafo.add_run(descricao)
    elif not inserida:
        texto = descricao or alt
        if texto:
            doc.add_paragraph(texto)
        if alt and descricao and alt not in descricao:
            from core.utils.logger import logger
            logger.warning(
                "Imagem nao embutida no DOCX; conteudo visual entregue "
                "apenas como texto"
            )


def _adicionar_caption_e_resumo(doc: Document, block: dict) -> None:
    caption = str(block.get("caption") or "").strip()
    if caption:
        paragrafo = doc.add_paragraph()
        paragrafo.add_run(caption).bold = True

    for chave, rotulo in (("x", "Eixo horizontal"), ("y", "Eixo vertical")):
        eixo = (block.get("axes") or {}).get(chave) or {}
        if not isinstance(eixo, dict) or not eixo.get("label"):
            continue
        partes = [str(eixo["label"])]
        if eixo.get("unit"):
            partes.append(f"em {eixo['unit']}")
        if eixo.get("min") is not None and eixo.get("max") is not None:
            partes.append(f"de {eixo['min']} a {eixo['max']}")
        if eixo.get("step") is not None:
            partes.append(f"intervalo {eixo['step']}")
        doc.add_paragraph(f"{rotulo}: {', '.join(partes)}.")

    resumo = str(block.get("summary") or "").strip()
    if resumo:
        doc.add_paragraph(resumo)


def _marcar_idioma_da_celula(celula, idioma: str, idioma_doc: str) -> None:
    try:
        from pipeline.localizacao import codigo_bcp47, precisa_marcar

        if not precisa_marcar(idioma, idioma_doc):
            return

        from docx.oxml.ns import qn

        codigo = codigo_bcp47(idioma)
        for paragrafo in celula.paragraphs:
            for run in paragrafo.runs:
                rPr = run._element.get_or_add_rPr()
                lang = rPr.find(qn("w:lang"))
                if lang is None:
                    lang = rPr.makeelement(qn("w:lang"), {})
                    rPr.append(lang)
                lang.set(qn("w:val"), codigo)
    except Exception:
        pass
